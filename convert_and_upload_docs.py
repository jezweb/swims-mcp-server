#!/usr/bin/env python3
"""
Convert DOCX/DOC files to PDF and upload to R2
"""

import os
import sys
from pathlib import Path
import boto3
from botocore.config import Config
import subprocess
import tempfile

# Try to import python-docx and reportlab for conversion
try:
    from docx import Document
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
    import io
    CONVERSION_AVAILABLE = True
except ImportError:
    CONVERSION_AVAILABLE = False
    print("Warning: python-docx and reportlab not available for conversion")

# R2 Configuration
R2_ACCOUNT_ID = os.getenv("R2_ACCOUNT_ID", "0460574641fdbb98159c98ebf593e2bd")
R2_ACCESS_KEY_ID = os.getenv("R2_ACCESS_KEY_ID")
R2_SECRET_ACCESS_KEY = os.getenv("R2_SECRET_ACCESS_KEY")
R2_BUCKET_NAME = "swms-regulations"

def convert_docx_to_pdf(docx_path):
    """Convert DOCX to PDF using python-docx and reportlab"""
    if not CONVERSION_AVAILABLE:
        return None
    
    try:
        # Read DOCX
        doc = Document(docx_path)
        
        # Create PDF in memory
        pdf_buffer = io.BytesIO()
        pdf = SimpleDocTemplate(pdf_buffer, pagesize=letter,
                               rightMargin=72, leftMargin=72,
                               topMargin=72, bottomMargin=18)
        
        # Container for the 'Flowable' objects
        story = []
        styles = getSampleStyleSheet()
        
        # Create custom styles
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontSize=14,
            textColor=colors.HexColor('#1a1a1a'),
            spaceAfter=12,
            fontName='Helvetica-Bold'
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor('#333333'),
            spaceAfter=6,
            fontName='Helvetica'
        )
        
        # Process paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            # Clean text for reportlab
            text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            
            # Detect headings
            if para.style and 'Heading' in para.style.name:
                story.append(Paragraph(text, heading_style))
            elif para.runs and para.runs[0].bold:
                story.append(Paragraph(text, heading_style))
            else:
                story.append(Paragraph(text, normal_style))
            
            story.append(Spacer(1, 6))
        
        # Process tables if any
        for table in doc.tables:
            data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    cell_text = cell_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    row_data.append(Paragraph(cell_text, normal_style))
                data.append(row_data)
            
            if data:
                t = Table(data)
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                story.append(t)
                story.append(Spacer(1, 12))
        
        # Build PDF
        pdf.build(story)
        
        # Get PDF bytes
        pdf_bytes = pdf_buffer.getvalue()
        pdf_buffer.close()
        
        return pdf_bytes
        
    except Exception as e:
        print(f"Error converting {docx_path}: {e}")
        return None

def get_r2_client():
    """Create and return an R2 client"""
    if not R2_ACCESS_KEY_ID or not R2_SECRET_ACCESS_KEY:
        print("Error: R2_ACCESS_KEY_ID and R2_SECRET_ACCESS_KEY must be set")
        sys.exit(1)
    
    endpoint_url = f"https://{R2_ACCOUNT_ID}.r2.cloudflarestorage.com"
    
    return boto3.client(
        service_name='s3',
        endpoint_url=endpoint_url,
        aws_access_key_id=R2_ACCESS_KEY_ID,
        aws_secret_access_key=R2_SECRET_ACCESS_KEY,
        config=Config(
            signature_version='s3v4',
            s3={'addressing_style': 'path'}
        ),
        region_name='auto'
    )

def upload_to_r2(client, file_bytes, r2_key):
    """Upload file bytes to R2"""
    try:
        client.put_object(
            Bucket=R2_BUCKET_NAME,
            Key=r2_key,
            Body=file_bytes,
            ContentType='application/pdf'
        )
        print(f"  ✓ Uploaded: {r2_key}")
        return True
    except Exception as e:
        print(f"  ✗ Failed to upload {r2_key}: {e}")
        return False

def main():
    """Convert DOCX/DOC files to PDF and upload to R2"""
    print("Converting DOCX/DOC files to PDF and uploading to R2")
    print("=" * 60)
    
    # Get R2 client
    client = get_r2_client()
    
    # Find all DOCX/DOC files
    base_dir = Path("regulatory_documents")
    docx_files = list(base_dir.glob("**/*.docx"))
    doc_files = list(base_dir.glob("**/*.doc"))
    all_files = docx_files + doc_files
    
    if not all_files:
        print("No DOCX/DOC files found")
        return
    
    print(f"Found {len(all_files)} DOCX/DOC files to convert and upload\n")
    
    converted = 0
    uploaded = 0
    failed = []
    
    for doc_path in all_files:
        relative_path = doc_path.relative_to(base_dir)
        pdf_name = str(relative_path).rsplit('.', 1)[0] + '-converted.pdf'
        
        print(f"Converting: {relative_path}")
        
        # Convert to PDF
        pdf_bytes = convert_docx_to_pdf(doc_path)
        
        if pdf_bytes:
            converted += 1
            print(f"  ✓ Converted to PDF ({len(pdf_bytes):,} bytes)")
            
            # Upload to R2
            if upload_to_r2(client, pdf_bytes, pdf_name):
                uploaded += 1
            else:
                failed.append(pdf_name)
        else:
            print(f"  ✗ Failed to convert")
            failed.append(str(relative_path))
    
    print("\n" + "=" * 60)
    print(f"Conversion Summary:")
    print(f"  Total files: {len(all_files)}")
    print(f"  Converted: {converted}")
    print(f"  Uploaded: {uploaded}")
    print(f"  Failed: {len(failed)}")
    
    if failed:
        print("\nFailed files:")
        for f in failed:
            print(f"  - {f}")
    
    # List all objects in bucket
    print("\n" + "=" * 60)
    print("Current bucket contents:")
    try:
        response = client.list_objects_v2(Bucket=R2_BUCKET_NAME)
        if 'Contents' in response:
            pdf_count = 0
            for obj in response['Contents']:
                if obj['Key'].endswith('.pdf'):
                    pdf_count += 1
            print(f"Total PDF files in bucket: {pdf_count}")
        else:
            print("Bucket is empty")
    except Exception as e:
        print(f"Could not list bucket contents: {e}")
    
    print("\n✓ Conversion and upload complete!")

if __name__ == "__main__":
    main()