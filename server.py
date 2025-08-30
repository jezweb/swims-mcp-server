"""
SWMS MCP Server - Safe Work Method Statement Analysis
"""

import os
import io
import json
import base64
import tempfile
import requests
from pathlib import Path
from typing import Dict, Any, List, Optional
from fastmcp import FastMCP
from dotenv import load_dotenv
from google import genai
from google.genai import types

# Import R2 context manager
from r2_context import R2ContextManager

# Import libraries for DOCX to PDF conversion
try:
    from docx import Document
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
    DOCX_CONVERSION_AVAILABLE = True
except ImportError:
    DOCX_CONVERSION_AVAILABLE = False

# Load environment variables
load_dotenv()

# Configure Gemini API client
api_key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
client = None
r2_context = None
if api_key:
    client = genai.Client(api_key=api_key)
    r2_context = R2ContextManager(client)

# MUST be at module level for FastMCP Cloud
mcp = FastMCP("swms-analysis-server")

def convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    """
    Convert DOCX bytes to PDF bytes using python-docx and reportlab.
    This is a simplified conversion that preserves text and basic formatting.
    """
    if not DOCX_CONVERSION_AVAILABLE:
        raise ImportError("DOCX conversion libraries not available")
    
    # Read DOCX document
    doc = Document(io.BytesIO(docx_bytes))
    
    # Create PDF in memory
    pdf_buffer = io.BytesIO()
    pdf = SimpleDocTemplate(pdf_buffer, pagesize=letter,
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)
    
    # Container for the 'Flowable' objects
    story = []
    styles = getSampleStyleSheet()
    
    # Create custom styles for better formatting
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading1'],
        fontSize=14,
        textColor=colors.HexColor('#1a1a1a'),
        spaceAfter=12,
        fontName='Helvetica-Bold'
    )
    
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors.HexColor('#333333'),
        spaceAfter=6,
        fontName='Helvetica'
    )
    
    # Process paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        # Clean text for reportlab (escape special characters)
        text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        
        # Detect headings based on style or formatting
        if para.style and 'Heading' in para.style.name:
            story.append(Paragraph(text, heading_style))
        elif para.runs and para.runs[0].bold:
            story.append(Paragraph(text, heading_style))
        else:
            story.append(Paragraph(text, normal_style))
        
        story.append(Spacer(1, 6))
    
    # Process tables if any
    for table in doc.tables:
        data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                # Clean cell text
                cell_text = cell.text.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                row_data.append(cell_text)
            data.append(row_data)
        
        if data:
            # Create table with basic styling
            t = Table(data)
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
            ]))
            story.append(t)
            story.append(Spacer(1, 12))
    
    # Build PDF
    pdf.build(story)
    
    # Get PDF bytes
    pdf_bytes = pdf_buffer.getvalue()
    pdf_buffer.close()
    
    return pdf_bytes

@mcp.tool()
async def upload_swms_document(
    file_content: str,
    file_name: str,
    mime_type: Optional[str] = None
) -> Dict[str, Any]:
    """
    Upload a SWMS document from base64-encoded content to Gemini API.
    
    Args:
        file_content: Base64-encoded file content
        file_name: Name of the file (e.g., "safety_plan.pdf")
        mime_type: Optional MIME type (auto-detected if not provided)
        
    Returns:
        Dictionary with upload status and document ID
    """
    try:
        if not client:
            return {
                "status": "error",
                "message": "Gemini API key not configured"
            }
        
        # Decode base64 content
        try:
            file_bytes = base64.b64decode(file_content)
        except Exception as e:
            return {
                "status": "error",
                "message": f"Invalid base64 content: {str(e)}"
            }
        
        # Auto-detect MIME type if not provided
        converted_from_docx = False
        original_file_name = file_name
        
        if not mime_type:
            extension = Path(file_name).suffix.lower()
            if extension == '.pdf':
                mime_type = 'application/pdf'
            elif extension in ['.docx', '.doc']:
                mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            else:
                return {
                    "status": "error",
                    "message": f"Unsupported file format: {extension}. Only PDF and DOCX are supported."
                }
        
        # Convert DOCX to PDF if needed
        if mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or \
           file_name.lower().endswith('.docx'):
            if DOCX_CONVERSION_AVAILABLE:
                try:
                    # Convert DOCX to PDF
                    file_bytes = convert_docx_to_pdf(file_bytes)
                    # Update file name and mime type
                    file_name = Path(file_name).stem + '.pdf'
                    mime_type = 'application/pdf'
                    converted_from_docx = True
                except Exception as e:
                    return {
                        "status": "error",
                        "message": f"Failed to convert DOCX to PDF: {str(e)}"
                    }
            else:
                return {
                    "status": "error",
                    "message": "DOCX files require conversion to PDF, but conversion libraries are not available"
                }
        
        # Create temporary file for upload
        with tempfile.NamedTemporaryFile(suffix='.pdf' if converted_from_docx else Path(file_name).suffix, delete=False) as temp_file:
            temp_file.write(file_bytes)
            temp_path = temp_file.name
        
        try:
            # Upload file to Gemini
            uploaded_file = client.files.upload(
                file=temp_path,
                config=types.UploadFileConfig(
                    display_name=file_name,
                    mime_type=mime_type
                )
            )
            
            response = {
                "status": "success",
                "message": f"Document {file_name} uploaded successfully",
                "document_id": uploaded_file.name,
                "file_info": {
                    "name": uploaded_file.display_name,
                    "mime_type": uploaded_file.mime_type,
                    "uri": uploaded_file.uri,
                    "size_bytes": len(file_bytes)
                }
            }
            
            if converted_from_docx:
                response["conversion_info"] = {
                    "original_format": "DOCX",
                    "original_name": original_file_name,
                    "converted_to": "PDF",
                    "note": "Document was automatically converted from DOCX to PDF for Gemini compatibility"
                }
            
            return response
        finally:
            # Clean up temporary file
            os.unlink(temp_path)
            
    except Exception as e:
        return {
            "status": "error", 
            "message": f"Failed to upload document: {str(e)}"
        }

@mcp.tool()
async def upload_swms_from_url(url: str) -> Dict[str, Any]:
    """
    Upload a SWMS document from a URL to Gemini API.
    
    Args:
        url: URL of the SWMS document (PDF or DOCX)
        
    Returns:
        Dictionary with upload status and document ID
    """
    try:
        if not client:
            return {
                "status": "error",
                "message": "Gemini API key not configured"
            }
        
        # Download file from URL
        try:
            response = requests.get(url, timeout=30)
            response.raise_for_status()
        except requests.RequestException as e:
            return {
                "status": "error",
                "message": f"Failed to download from URL: {str(e)}"
            }
        
        # Determine file name and MIME type from URL and headers
        file_name = url.split('/')[-1].split('?')[0]
        if not file_name or '.' not in file_name:
            file_name = "document.pdf"  # Default name
        
        converted_from_docx = False
        original_file_name = file_name
        file_bytes = response.content
        
        content_type = response.headers.get('content-type', '')
        if 'pdf' in content_type:
            mime_type = 'application/pdf'
            if not file_name.endswith('.pdf'):
                file_name = file_name.split('.')[0] + '.pdf'
        elif 'wordprocessingml' in content_type or 'msword' in content_type:
            mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            if not file_name.endswith('.docx'):
                file_name = file_name.split('.')[0] + '.docx'
        else:
            # Try to detect from file extension
            extension = Path(file_name).suffix.lower()
            if extension == '.pdf':
                mime_type = 'application/pdf'
            elif extension in ['.docx', '.doc']:
                mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            else:
                return {
                    "status": "error",
                    "message": f"Could not determine file type from URL. Ensure it's a PDF or DOCX file."
                }
        
        # Convert DOCX to PDF if needed
        if mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            if DOCX_CONVERSION_AVAILABLE:
                try:
                    # Convert DOCX to PDF
                    file_bytes = convert_docx_to_pdf(file_bytes)
                    # Update file name and mime type
                    file_name = Path(file_name).stem + '.pdf'
                    mime_type = 'application/pdf'
                    converted_from_docx = True
                except Exception as e:
                    return {
                        "status": "error",
                        "message": f"Failed to convert DOCX to PDF: {str(e)}"
                    }
            else:
                return {
                    "status": "error",
                    "message": "DOCX files require conversion to PDF, but conversion libraries are not available"
                }
        
        # Create temporary file for upload
        with tempfile.NamedTemporaryFile(suffix='.pdf' if converted_from_docx else Path(file_name).suffix, delete=False) as temp_file:
            temp_file.write(file_bytes)
            temp_path = temp_file.name
        
        try:
            # Upload file to Gemini
            uploaded_file = client.files.upload(
                file=temp_path,
                config=types.UploadFileConfig(
                    display_name=file_name,
                    mime_type=mime_type
                )
            )
            
            response_data = {
                "status": "success",
                "message": f"Document {file_name} uploaded successfully from URL",
                "document_id": uploaded_file.name,
                "file_info": {
                    "name": uploaded_file.display_name,
                    "mime_type": uploaded_file.mime_type,
                    "uri": uploaded_file.uri,
                    "size_bytes": len(file_bytes),
                    "source_url": url
                }
            }
            
            if converted_from_docx:
                response_data["conversion_info"] = {
                    "original_format": "DOCX",
                    "original_name": original_file_name,
                    "converted_to": "PDF",
                    "note": "Document was automatically converted from DOCX to PDF for Gemini compatibility"
                }
            
            return response_data
        finally:
            # Clean up temporary file
            os.unlink(temp_path)
            
    except Exception as e:
        return {
            "status": "error", 
            "message": f"Failed to upload document from URL: {str(e)}"
        }

@mcp.tool()
async def analyze_swms_compliance(
    document_id: str,
    jurisdiction: Optional[str] = "nsw"
) -> Dict[str, Any]:
    """
    Analyze a SWMS document for WHS compliance using Gemini API.
    
    Args:
        document_id: ID of the uploaded SWMS document
        jurisdiction: State/territory jurisdiction (nsw, vic, qld, wa, sa, tas, act, nt, national)
        
    Returns:
        Detailed compliance assessment report
    """
    try:
        if not client:
            return {
                "status": "error",
                "message": "Gemini API key not configured"
            }
        
        # Get the uploaded file
        try:
            uploaded_file = client.files.get(name=document_id)
        except Exception as e:
            return {
                "status": "error",
                "message": f"Document not found: {document_id}"
            }
        
        # Get jurisdiction-specific context if R2 context manager is available
        context_files = []
        jurisdiction_info = {}
        if r2_context and jurisdiction:
            try:
                # Get regulatory document file IDs from R2
                context_files = r2_context.get_context_files(jurisdiction)
                # Get jurisdiction-specific information
                jurisdiction_info = r2_context.get_jurisdiction_context(jurisdiction)
            except Exception as e:
                print(f"Warning: Could not load R2 context: {e}")
        
        # Build jurisdiction-aware prompt
        jurisdiction_upper = jurisdiction.upper() if jurisdiction else "NSW"
        legislation = jurisdiction_info.get("legislation", "Work Health and Safety Regulation 2017")
        regulator = jurisdiction_info.get("regulatory_body", "SafeWork NSW")
        
        # Adjust terminology for Victoria
        terminology = "WHS" if jurisdiction != "vic" else "OHS"
        
        # SWMS Assessment Prompt with jurisdiction awareness
        assessment_prompt = f"""
## System Prompt for Assessing Safe Work Method Statements (SWMS) in {jurisdiction_upper} Construction

**Objective:** Analyze the provided Safe Work Method Statement (SWMS) for completeness and compliance with {legislation}. Generate a detailed compliance report.

**Jurisdiction:** {jurisdiction_upper}
**Regulatory Body:** {regulator}
**Legislation Framework:** {legislation}

**Instructions:**

You are an AI assistant with expertise in {jurisdiction_upper} Work Health and Safety ({terminology}) legislation for the construction industry. Assess the SWMS document according to these key areas:

**1. Document Control and Administrative Compliance:**
- Project details (name, address, Principal Contractor, subcontractor, ABN)
- Version control, document numbers, dates
- Personnel identification and responsibilities
- Worker sign-off provisions

**2. Identification of High-Risk Construction Work (HRCW):**
- Explicit identification of HRCW activities
- Correct categorization per NSW WHS Regulation 2017 (18 categories)

**3. Hazard Identification and Risk Assessment:**
- Logical task breakdown
- Site-specific (not generic) hazard identification
- Clear risk descriptions for each hazard

**4. Control Measures:**
- Adherence to hierarchy of controls (Elimination → Substitution → Isolation → Engineering → Administrative → PPE)
- Sufficient detail for worker understanding
- Clear implementation descriptions

**5. Monitoring, Review, and Communication:**
- Defined monitoring responsibilities
- Review triggers (work changes, incidents, ineffective controls)
- Communication plans for workers

**6. Consultation:**
- Evidence of worker consultation
- Sign-off sheets or consultation records

**Output Format Required:**

Return a JSON object with this exact structure:

{
  "status": "success",
  "project_details": {
    "project_name": "[extracted or 'Not specified']",
    "principal_contractor": "[extracted or 'Not specified']",
    "subcontractor": "[extracted or 'Not specified']",
    "swms_title": "[extracted or 'Not specified']"
  },
  "overall_assessment": "Compliant|Partially Compliant|Non-Compliant",
  "summary": "[Brief high-level summary of compliance status]",
  "detailed_analysis": {
    "document_control": {
      "status": "Compliant|Partially Compliant|Non-Compliant",
      "comments": "[Specific findings and recommendations]"
    },
    "hrcw_identification": {
      "status": "Compliant|Partially Compliant|Non-Compliant",
      "comments": "[Specific findings and recommendations]"
    },
    "hazard_identification": {
      "status": "Compliant|Partially Compliant|Non-Compliant",
      "comments": "[Specific findings and recommendations]"
    },
    "control_measures": {
      "status": "Compliant|Partially Compliant|Non-Compliant",
      "comments": "[Specific findings and recommendations]"
    },
    "monitoring_review": {
      "status": "Compliant|Partially Compliant|Non-Compliant",
      "comments": "[Specific findings and recommendations]"
    },
    "consultation": {
      "status": "Compliant|Partially Compliant|Non-Compliant",
      "comments": "[Specific findings and recommendations]"
    }
  },
  "urgent_actions": [
    "[List critical non-compliances that must be addressed before work commences]"
  ],
  "recommendations": [
    "[List other improvements for full compliance]"
  ]
}

Analyze the attached SWMS document thoroughly and provide the assessment in the exact JSON format above.
"""
        
        # Build contents list with prompt and document
        contents = [assessment_prompt]
        
        # Add context files if available
        if context_files:
            for file_id in context_files:
                try:
                    context_file = client.files.get(name=file_id)
                    contents.append(
                        types.Part.from_uri(
                            file_uri=context_file.uri,
                            mime_type=context_file.mime_type
                        )
                    )
                except Exception as e:
                    print(f"Warning: Could not add context file {file_id}: {e}")
        
        # Add the main SWMS document to analyze
        contents.append(
            types.Part.from_uri(
                file_uri=uploaded_file.uri,
                mime_type=uploaded_file.mime_type
            )
        )
        
        # Generate analysis using Gemini model
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=contents
        )
        
        # Parse the JSON response
        try:
            # Extract JSON from response text (handle potential markdown formatting)
            response_text = response.text.strip()
            if response_text.startswith('```json'):
                response_text = response_text[7:-3].strip()
            elif response_text.startswith('```'):
                response_text = response_text[3:-3].strip()
            
            analysis_result = json.loads(response_text)
            
            # Ensure required structure
            if "status" not in analysis_result:
                analysis_result["status"] = "success"
                
            return analysis_result
            
        except json.JSONDecodeError as e:
            # Fallback if JSON parsing fails
            return {
                "status": "success",
                "overall_assessment": "Analysis Completed",
                "summary": "Document analyzed but response format needs adjustment",
                "raw_response": response.text[:2000],  # Truncate for safety
                "parse_error": str(e)
            }
        
    except Exception as e:
        return {
            "status": "error",
            "message": f"Failed to analyze document: {str(e)}"
        }

@mcp.tool()
async def analyze_swms_text(
    document_text: str,
    document_name: Optional[str] = "SWMS Document",
    jurisdiction: Optional[str] = "nsw"
) -> Dict[str, Any]:
    """
    Analyze SWMS text content directly for WHS compliance.
    
    Args:
        document_text: The SWMS document content as plain text or markdown
        document_name: Optional name for the document
        jurisdiction: State/territory jurisdiction (nsw, vic, qld, wa, sa, tas, act, nt, national)
        
    Returns:
        Detailed compliance assessment report
    """
    try:
        if not client:
            return {
                "status": "error",
                "message": "Gemini API key not configured"
            }
        
        # NSW SWMS Assessment Prompt (same as file-based analysis)
        assessment_prompt = """
## System Prompt for Assessing Safe Work Method Statements (SWMS) in NSW Construction

**Objective:** Analyze the provided Safe Work Method Statement (SWMS) for completeness and compliance with NSW Work Health and Safety (WHS) Regulation 2017. Generate a detailed compliance report.

**Instructions:**

You are an AI assistant with expertise in NSW Work Health and Safety (WHS) legislation for the construction industry. Assess the SWMS document according to these key areas:

**1. Document Control and Administrative Compliance:**
- Project details (name, address, Principal Contractor, subcontractor, ABN)
- Version control, document numbers, dates
- Personnel identification and responsibilities
- Worker sign-off provisions

**2. Identification of High-Risk Construction Work (HRCW):**
- Explicit identification of HRCW activities
- Correct categorization per NSW WHS Regulation 2017 (18 categories)

**3. Hazard Identification and Risk Assessment:**
- Logical task breakdown
- Site-specific (not generic) hazard identification
- Clear risk descriptions for each hazard

**4. Control Measures:**
- Adherence to hierarchy of controls (Elimination → Substitution → Isolation → Engineering → Administrative → PPE)
- Sufficient detail for worker understanding
- Clear implementation descriptions

**5. Monitoring, Review, and Communication:**
- Defined monitoring responsibilities
- Review triggers (work changes, incidents, ineffective controls)
- Communication plans for workers

**6. Consultation:**
- Evidence of worker consultation
- Sign-off sheets or consultation records

**Output Format Required:**

Return a JSON object with this exact structure:

{
  "status": "success",
  "project_details": {
    "project_name": "[extracted or 'Not specified']",
    "principal_contractor": "[extracted or 'Not specified']",
    "subcontractor": "[extracted or 'Not specified']",
    "swms_title": "[extracted or 'Not specified']"
  },
  "overall_assessment": "Compliant|Partially Compliant|Non-Compliant",
  "summary": "[Brief high-level summary of compliance status]",
  "detailed_analysis": {
    "document_control": {
      "status": "Compliant|Partially Compliant|Non-Compliant",
      "comments": "[Specific findings and recommendations]"
    },
    "hrcw_identification": {
      "status": "Compliant|Partially Compliant|Non-Compliant",
      "comments": "[Specific findings and recommendations]"
    },
    "hazard_identification": {
      "status": "Compliant|Partially Compliant|Non-Compliant",
      "comments": "[Specific findings and recommendations]"
    },
    "control_measures": {
      "status": "Compliant|Partially Compliant|Non-Compliant",
      "comments": "[Specific findings and recommendations]"
    },
    "monitoring_review": {
      "status": "Compliant|Partially Compliant|Non-Compliant",
      "comments": "[Specific findings and recommendations]"
    },
    "consultation": {
      "status": "Compliant|Partially Compliant|Non-Compliant",
      "comments": "[Specific findings and recommendations]"
    }
  },
  "urgent_actions": [
    "[List critical non-compliances that must be addressed before work commences]"
  ],
  "recommendations": [
    "[List other improvements for full compliance]"
  ]
}

Analyze the attached SWMS document text thoroughly and provide the assessment in the exact JSON format above.

DOCUMENT NAME: """ + document_name + """

DOCUMENT CONTENT:
""" + document_text
        
        # Generate analysis using Gemini model
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=assessment_prompt
        )
        
        # Parse the JSON response
        try:
            # Extract JSON from response text (handle potential markdown formatting)
            response_text = response.text.strip()
            if response_text.startswith('```json'):
                response_text = response_text[7:-3].strip()
            elif response_text.startswith('```'):
                response_text = response_text[3:-3].strip()
            
            analysis_result = json.loads(response_text)
            
            # Ensure required structure
            if "status" not in analysis_result:
                analysis_result["status"] = "success"
                
            return analysis_result
            
        except json.JSONDecodeError as e:
            # Fallback if JSON parsing fails
            return {
                "status": "success",
                "overall_assessment": "Analysis Completed",
                "summary": "Document analyzed but response format needs adjustment",
                "raw_response": response.text[:2000],  # Truncate for safety
                "parse_error": str(e)
            }
        
    except Exception as e:
        return {
            "status": "error",
            "message": f"Failed to analyze document: {str(e)}"
        }

@mcp.tool()
async def analyze_swms_custom(
    document_id: str,
    analysis_prompt: str,
    output_format: Optional[str] = "json"
) -> Dict[str, Any]:
    """
    Perform custom analysis on a SWMS document using a user-provided prompt.
    
    Args:
        document_id: ID of the uploaded SWMS document
        analysis_prompt: Custom prompt for the analysis
        output_format: Expected output format ('json', 'text', or 'structured')
        
    Returns:
        Analysis results based on the custom prompt
    """
    try:
        if not client:
            return {
                "status": "error",
                "message": "Gemini API key not configured"
            }
        
        # Get the uploaded file
        try:
            uploaded_file = client.files.get(name=document_id)
        except Exception as e:
            return {
                "status": "error",
                "message": f"Document not found: {document_id}"
            }
        
        # Add format instructions based on output_format
        format_instructions = ""
        if output_format == "json":
            format_instructions = "\n\nPlease return your response as valid JSON."
        elif output_format == "structured":
            format_instructions = "\n\nPlease structure your response with clear headings and bullet points."
        else:  # text
            format_instructions = "\n\nPlease provide a clear, detailed text response."
        
        full_prompt = analysis_prompt + format_instructions
        
        # Generate analysis using Gemini model
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=[
                full_prompt,
                types.Part.from_uri(
                    file_uri=uploaded_file.uri,
                    mime_type=uploaded_file.mime_type
                )
            ]
        )
        
        # Handle response based on expected format
        if output_format == "json":
            try:
                # Try to parse as JSON
                response_text = response.text.strip()
                if response_text.startswith('```json'):
                    response_text = response_text[7:-3].strip()
                elif response_text.startswith('```'):
                    response_text = response_text[3:-3].strip()
                
                result = json.loads(response_text)
                return {
                    "status": "success",
                    "output_format": "json",
                    "result": result
                }
            except json.JSONDecodeError:
                # Fall back to text if JSON parsing fails
                return {
                    "status": "success",
                    "output_format": "text",
                    "result": response.text,
                    "note": "JSON parsing failed, returning as text"
                }
        else:
            # Return as text for other formats
            return {
                "status": "success",
                "output_format": output_format,
                "result": response.text
            }
        
    except Exception as e:
        return {
            "status": "error",
            "message": f"Failed to perform custom analysis: {str(e)}"
        }

@mcp.tool()
async def get_compliance_score(
    document_id: str,
    weighted: bool = True,
    jurisdiction: Optional[str] = "nsw"
) -> Dict[str, Any]:
    """
    Calculate numerical compliance scores for a SWMS document.
    
    Args:
        document_id: ID of the uploaded SWMS document
        weighted: If True, weight scores by importance (HRCW and controls weighted higher)
        
    Returns:
        Compliance scores for each area and overall score
    """
    try:
        if not client:
            return {
                "status": "error",
                "message": "Gemini API key not configured"
            }
        
        # Get the uploaded file
        try:
            uploaded_file = client.files.get(name=document_id)
        except Exception as e:
            return {
                "status": "error",
                "message": f"Document not found: {document_id}"
            }
        
        # Scoring prompt
        scoring_prompt = """
Analyze this SWMS document and provide numerical compliance scores (0-100) for NSW WHS compliance.

Score each area based on:
- 0-25: Non-compliant (critical elements missing)
- 26-50: Partially compliant (significant gaps)
- 51-75: Mostly compliant (minor improvements needed)
- 76-100: Fully compliant (meets or exceeds requirements)

Return a JSON object with this exact structure:
{
  "scores": {
    "document_control": {
      "score": [0-100],
      "justification": "[Brief reason for score]"
    },
    "hrcw_identification": {
      "score": [0-100],
      "justification": "[Brief reason for score]"
    },
    "hazard_identification": {
      "score": [0-100],
      "justification": "[Brief reason for score]"
    },
    "control_measures": {
      "score": [0-100],
      "justification": "[Brief reason for score]"
    },
    "monitoring_review": {
      "score": [0-100],
      "justification": "[Brief reason for score]"
    },
    "consultation": {
      "score": [0-100],
      "justification": "[Brief reason for score]"
    }
  }
}
"""
        
        # Generate analysis using Gemini model
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=[
                scoring_prompt,
                types.Part.from_uri(
                    file_uri=uploaded_file.uri,
                    mime_type=uploaded_file.mime_type
                )
            ]
        )
        
        # Parse the JSON response
        try:
            response_text = response.text.strip()
            if response_text.startswith('```json'):
                response_text = response_text[7:-3].strip()
            elif response_text.startswith('```'):
                response_text = response_text[3:-3].strip()
            
            result = json.loads(response_text)
            scores = result.get("scores", {})
            
            # Calculate overall score
            if weighted:
                # Weighted importance (total = 100%)
                weights = {
                    "document_control": 0.10,
                    "hrcw_identification": 0.25,  # Critical for safety
                    "hazard_identification": 0.20,
                    "control_measures": 0.25,  # Critical for safety
                    "monitoring_review": 0.10,
                    "consultation": 0.10
                }
            else:
                # Equal weighting
                weights = {k: 1/6 for k in scores.keys()}
            
            overall_score = 0
            for area, data in scores.items():
                if isinstance(data, dict) and "score" in data:
                    overall_score += data["score"] * weights.get(area, 0)
            
            return {
                "status": "success",
                "overall_score": round(overall_score, 1),
                "weighted": weighted,
                "area_scores": scores,
                "weights_used": weights if weighted else "equal",
                "compliance_level": (
                    "Non-Compliant" if overall_score < 26 else
                    "Partially Compliant" if overall_score < 51 else
                    "Mostly Compliant" if overall_score < 76 else
                    "Fully Compliant"
                )
            }
            
        except (json.JSONDecodeError, KeyError) as e:
            return {
                "status": "error",
                "message": f"Failed to parse scoring response: {str(e)}",
                "raw_response": response.text[:1000]
            }
        
    except Exception as e:
        return {
            "status": "error",
            "message": f"Failed to calculate compliance score: {str(e)}"
        }

@mcp.tool()
async def quick_check_swms(
    document_id: str,
    check_type: str
) -> Dict[str, Any]:
    """
    Perform a quick focused check on specific aspects of a SWMS document.
    
    Args:
        document_id: ID of the uploaded SWMS document
        check_type: Type of check to perform. Options:
                   - 'hrcw': Check for High-Risk Construction Work identification
                   - 'ppe': Check PPE requirements
                   - 'emergency': Check emergency procedures
                   - 'signatures': Check for sign-off sections
                   - 'hierarchy': Check hierarchy of controls
                   - 'hazards': Quick hazard identification check
        
    Returns:
        Quick check results focused on the specific aspect
    """
    try:
        if not client:
            return {
                "status": "error",
                "message": "Gemini API key not configured"
            }
        
        # Get the uploaded file
        try:
            uploaded_file = client.files.get(name=document_id)
        except Exception as e:
            return {
                "status": "error",
                "message": f"Document not found: {document_id}"
            }
        
        # Define check-specific prompts
        check_prompts = {
            "hrcw": """
Check this SWMS for High-Risk Construction Work (HRCW) identification.
Look for activities from the 18 HRCW categories in NSW WHS Regulation 2017.
Return JSON: {"hrcw_found": [list of HRCW activities], "properly_identified": true/false, "missing": [any likely HRCW not explicitly identified]}
""",
            "ppe": """
Check PPE requirements in this SWMS.
Return JSON: {"ppe_specified": true/false, "ppe_items": [list of PPE required], "task_specific_ppe": {task: [ppe_items]}, "gaps": [missing PPE]}
""",
            "emergency": """
Check emergency procedures in this SWMS.
Return JSON: {"emergency_procedures": true/false, "contact_numbers": true/false, "evacuation_plan": true/false, "first_aid": true/false, "issues": [list of missing items]}
""",
            "signatures": """
Check for worker consultation and sign-off provisions.
Return JSON: {"sign_off_section": true/false, "consultation_evidence": true/false, "responsible_person": "name or not specified", "issues": [list of problems]}
""",
            "hierarchy": """
Check if control measures follow the hierarchy of controls.
Return JSON: {"hierarchy_followed": true/false, "elimination": [examples], "substitution": [examples], "engineering": [examples], "administrative": [examples], "ppe": [examples], "issues": [problems with hierarchy application]}
""",
            "hazards": """
Quick check of hazard identification completeness.
Return JSON: {"hazards_identified": [list of hazards], "site_specific": true/false, "generic_only": true/false, "missing_common": [likely missing hazards], "count": number}
"""
        }
        
        # Get the appropriate prompt
        prompt = check_prompts.get(check_type)
        if not prompt:
            return {
                "status": "error",
                "message": f"Invalid check_type: {check_type}. Valid options: {list(check_prompts.keys())}"
            }
        
        # Add instruction for clean JSON
        prompt += "\n\nReturn ONLY valid JSON, no markdown formatting or explanations."
        
        # Generate analysis using Gemini model
        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=[
                prompt,
                types.Part.from_uri(
                    file_uri=uploaded_file.uri,
                    mime_type=uploaded_file.mime_type
                )
            ]
        )
        
        # Parse the JSON response
        try:
            response_text = response.text.strip()
            if response_text.startswith('```json'):
                response_text = response_text[7:-3].strip()
            elif response_text.startswith('```'):
                response_text = response_text[3:-3].strip()
            
            result = json.loads(response_text)
            
            return {
                "status": "success",
                "check_type": check_type,
                "result": result,
                "quick_summary": _generate_quick_summary(check_type, result)
            }
            
        except json.JSONDecodeError as e:
            # Still return the text result if JSON parsing fails
            return {
                "status": "success",
                "check_type": check_type,
                "result": {"raw_response": response.text},
                "note": "Response was not valid JSON"
            }
        
    except Exception as e:
        return {
            "status": "error",
            "message": f"Failed to perform quick check: {str(e)}"
        }

def _generate_quick_summary(check_type: str, result: Dict) -> str:
    """Generate a quick text summary based on check results."""
    summaries = {
        "hrcw": lambda r: f"Found {len(r.get('hrcw_found', []))} HRCW activities. " +
                         ("Properly identified." if r.get('properly_identified') else "Issues with identification."),
        "ppe": lambda r: "PPE specified." if r.get('ppe_specified') else "PPE requirements missing or incomplete.",
        "emergency": lambda r: "Emergency procedures present." if r.get('emergency_procedures') else "Emergency procedures missing.",
        "signatures": lambda r: "Sign-off section present." if r.get('sign_off_section') else "Sign-off section missing.",
        "hierarchy": lambda r: "Hierarchy of controls followed." if r.get('hierarchy_followed') else "Issues with hierarchy of controls.",
        "hazards": lambda r: f"Identified {r.get('count', 0)} hazards. " +
                            ("Site-specific." if r.get('site_specific') else "Generic hazards only.")
    }
    
    summary_func = summaries.get(check_type)
    if summary_func and isinstance(result, dict):
        try:
            return summary_func(result)
        except:
            return "Check completed."
    return "Check completed."

@mcp.tool()
async def list_jurisdictions() -> Dict[str, Any]:
    """
    List all supported jurisdictions for SWMS analysis.
    
    Returns:
        Dictionary with supported jurisdictions and their details
    """
    jurisdictions = {
        "national": {
            "name": "National (Model Laws)",
            "legislation": "Model Work Health and Safety Act and Regulations",
            "regulator": "Safe Work Australia",
            "terminology": "WHS"
        },
        "nsw": {
            "name": "New South Wales",
            "legislation": "Work Health and Safety Act 2011 (NSW) and Regulation 2017 (NSW)",
            "regulator": "SafeWork NSW",
            "terminology": "WHS"
        },
        "vic": {
            "name": "Victoria",
            "legislation": "Occupational Health and Safety Act 2004 (Vic) and Regulations 2017 (Vic)",
            "regulator": "WorkSafe Victoria",
            "terminology": "OHS",
            "note": "Victoria uses OHS terminology and has not adopted model WHS laws"
        },
        "qld": {
            "name": "Queensland",
            "legislation": "Work Health and Safety Act 2011 (Qld) and Regulation 2011 (Qld)",
            "regulator": "Workplace Health and Safety Queensland",
            "terminology": "WHS"
        },
        "wa": {
            "name": "Western Australia",
            "legislation": "Work Health and Safety Act 2020 (WA) and Regulations 2022 (WA)",
            "regulator": "WorkSafe Western Australia",
            "terminology": "WHS",
            "note": "Adopted WHS laws in 2022 with some variations"
        },
        "sa": {
            "name": "South Australia",
            "legislation": "Work Health and Safety Act 2012 (SA) and Regulations 2012 (SA)",
            "regulator": "SafeWork SA",
            "terminology": "WHS"
        },
        "tas": {
            "name": "Tasmania",
            "legislation": "Work Health and Safety Act 2012 (Tas) and Regulations 2012 (Tas)",
            "regulator": "WorkSafe Tasmania",
            "terminology": "WHS"
        },
        "act": {
            "name": "Australian Capital Territory",
            "legislation": "Work Health and Safety Act 2011 (ACT) and Regulation 2011 (ACT)",
            "regulator": "WorkSafe ACT",
            "terminology": "WHS"
        },
        "nt": {
            "name": "Northern Territory",
            "legislation": "Work Health and Safety (National Uniform Legislation) Act 2011 (NT)",
            "regulator": "NT WorkSafe",
            "terminology": "WHS"
        }
    }
    
    return {
        "status": "success",
        "jurisdictions": jurisdictions,
        "default": "nsw",
        "note": "All analysis functions support jurisdiction parameter. R2 context documents will be included if available."
    }

@mcp.tool()
async def get_server_status() -> Dict[str, Any]:
    """
    Get the current status of the SWMS analysis server.
    
    Returns:
        Server status and configuration info
    """
    api_configured = bool(api_key)
    
    # Test API connection if configured
    api_status = "not_configured"
    if api_configured and client:
        try:
            # Simple test to verify API connectivity
            models = client.models.list()
            api_status = "active" if models else "connection_failed"
        except Exception as e:
            api_status = f"error: {str(e)}"
    
    return {
        "status": "active",
        "server_name": "SWMS Analysis Server",
        "version": "1.0.0",
        "supported_formats": ["PDF", "DOCX"],
        "gemini_api": {
            "configured": api_configured,
            "status": api_status
        },
        "capabilities": [
            "upload_swms_document",
            "upload_swms_from_url",
            "analyze_swms_text",
            "analyze_swms_compliance", 
            "get_server_status"
        ]
    }

# Remove the if __name__ == "__main__" block for FastMCP Cloud compatibility
# The server should be run using: fastmcp run server.py